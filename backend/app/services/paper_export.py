"""Paper export to .docx (question paper, answer key, response-sheet template).

Uses python-docx. The three variants cover the examiner workflow:
``paper`` renders the item stems, ``key`` repeats them with the correct answer,
and ``template`` produces a blank response sheet ready to be distributed.
"""

import io

from docx import Document
from docx.shared import Pt

from app.models import ExamPaper


def export_paper(paper: ExamPaper, variant: str) -> bytes:
    doc = Document()
    doc.add_heading(paper.title, level=0)

    ordered = [q for q in paper.questions]

    if variant == "template":
        _render_template(doc, ordered)
    elif variant == "key":
        _render_items(doc, ordered, with_key=True)
    else:
        _render_items(doc, ordered, with_key=False)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_items(doc: Document, ordered, with_key: bool) -> None:
    doc.add_heading("Answer Key" if with_key else "Question Paper", level=1)
    for i, link in enumerate(ordered, start=1):
        q = link.question
        doc.add_heading(f"{i}.", level=2)
        doc.add_paragraph(q.stem)
        for opt in q.options:
            doc.add_paragraph(f"{chr(ord('A') + opt.position)}. {opt.text}", style="List Bullet")
        if with_key:
            key = next((o for o in q.options if o.is_correct), None)
            if key:
                p = doc.add_paragraph()
                run = p.add_run(f"Answer: {chr(ord('A') + key.position)}")
                run.bold = True


def _render_template(doc: Document, ordered) -> None:
    doc.add_heading("Response Sheet", level=1)
    table = doc.add_table(rows=1, cols=len(ordered) + 2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Candidate"
    hdr[1].text = "ID"
    for j, link in enumerate(ordered, start=2):
        hdr[j].text = str(link.position)
    for _ in range(10):
        cells = table.add_row().cells
        cells[0].text = ""
        cells[1].text = ""
        for j in range(2, len(ordered) + 2):
            cells[j].text = "A  B  C  D"
